#!/usr/bin/env python3
"""Batch-apply text replacements as Word tracked changes + AI comments.

Two-stage pipeline:
  Stage 1 (docx-revisions): execute red delete / blue insert revisions
  Stage 2 (python-docx): attach explanatory comments to inserted runs
"""
import argparse
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.text.run import Run
from docx_revisions import RevisionDocument


def has_revision_markup(path: str) -> bool:
    with zipfile.ZipFile(path, "r") as z:
        with z.open("word/document.xml") as f:
            content = f.read().decode("utf-8")
            return "<w:ins" in content or "<w:del" in content


def has_comments_part(path: str) -> bool:
    doc = Document(path)
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            return True
    return False


def stage1_revisions(
    input_path: str,
    output_path: str,
    replacements: list[dict],
    author: str,
) -> dict:
    """Execute tracked replacements."""
    rdoc = RevisionDocument(input_path)
    total = 0
    failed = []
    for item in replacements:
        idx = item["paragraph_index"]
        search = item["search"]
        replace = item["replace"]
        para = rdoc.paragraphs[idx]
        count = para.replace_tracked(search, replace, author=author)
        total += count
        if count < 1:
            failed.append({"index": idx, "search_prefix": search[:60]})
    rdoc.save(output_path)
    return {"total": total, "failed": failed, "output": output_path}


def stage2_comments(
    input_path: str,
    output_path: str,
    replacements: list[dict],
    author: str,
    initials: str = "AI",
) -> dict:
    """Attach comments to inserted runs."""
    doc = Document(input_path)
    comment_map = {
        item["replace"][:60].strip(): item["comment"]
        for item in replacements
        if item.get("comment")
    }
    comment_count = 0
    for para in doc.paragraphs:
        for child in para._p:
            if child.tag == qn("w:ins"):
                ins_text = "".join(
                    t.text
                    for r in child
                    if r.tag == qn("w:r")
                    for t in r
                    if t.tag == qn("w:t")
                )
                matched = None
                for key, comment_text in comment_map.items():
                    if key in ins_text or ins_text[:60] in key:
                        matched = comment_text
                        break
                if matched:
                    for r in child:
                        if r.tag == qn("w:r"):
                            run_obj = Run(r, para)
                            doc.add_comment(
                                runs=[run_obj],
                                text=matched,
                                author=author,
                                initials=initials,
                            )
                            comment_count += 1
                            break
                    break
    doc.save(output_path)
    return {"comments": comment_count, "output": output_path}


def verify(output_path: str) -> dict:
    return {
        "has_revisions": has_revision_markup(output_path),
        "has_comments": has_comments_part(output_path),
    }


def main():
    parser = argparse.ArgumentParser(description="Apply diff replacements as Word tracked changes")
    parser.add_argument("--input", required=True, help="Original docx path")
    parser.add_argument("--replacements", required=True, help="JSON file with replacements")
    parser.add_argument("--output", required=True, help="Output docx path")
    parser.add_argument("--author", default="AI审阅助手", help="Revision author name")
    parser.add_argument("--initials", default="AI", help="Comment author initials")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: input not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    with open(args.replacements, "r", encoding="utf-8") as f:
        replacements = json.load(f)

    output_dir = Path(args.output).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    stage1_path = output_dir / "_stage1_revisions.docx"
    final_path = Path(args.output)

    print("=" * 50)
    print("Stage 1: Tracked replacements")
    print("=" * 50)
    r1 = stage1_revisions(str(input_path), str(stage1_path), replacements, args.author)
    print(f"Total replacements: {r1['total']}/{len(replacements)}")
    if r1["failed"]:
        print(f"Failed: {r1['failed']}")

    print()
    print("=" * 50)
    print("Stage 2: Attach comments")
    print("=" * 50)
    r2 = stage2_comments(str(stage1_path), str(final_path), replacements, args.author, args.initials)
    print(f"Comments added: {r2['comments']}")

    print()
    print("=" * 50)
    print("Verification")
    print("=" * 50)
    v = verify(str(final_path))
    print(f"Has revision markup: {v['has_revisions']}")
    print(f"Has comments part:   {v['has_comments']}")

    # Clean up stage1 temp file
    stage1_path.unlink(missing_ok=True)

    if not v["has_revisions"] or not v["has_comments"]:
        sys.exit(2)
    print(f"\nFinal output: {final_path}")


if __name__ == "__main__":
    main()
